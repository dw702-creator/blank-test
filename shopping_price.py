import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import random
from io import BytesIO
import nltk
import re

# ---------- NLTK: 필요한 데이터가 없을 때만 다운로드 ----------
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# ---------- Streamlit 페이지 설정 ----------
st.set_page_config(page_title="Your Blank Test Generator", layout="wide")
st.markdown(
    """
    <style>
    .stApp { font-family: "Segoe UI", Roboto, "Helvetica Neue", Arial; }
    .card {border-radius:10px; padding:16px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); background:#fff;}
    </style>
    """,
    unsafe_allow_html=True,
)

# Header
st.title("📝 Your Blank Test Generator")
st.markdown("학원용 스타일의 빈칸 시험지를 워드(.docx)로 자동 생성하고, 마지막 페이지에 답지도 포함합니다.")

# Sidebar for settings
with st.sidebar:
    st.header("설정")
    blank_ratio = st.slider("빈칸 비율 (%)", min_value=5, max_value=80, value=25, step=5,
                            help="문서 전체 단어 중 몇 %를 빈칸으로 만들지 설정합니다.")
    keep_punct = st.checkbox("구두점(.,?!)은 빈칸으로 선택하지 않음", value=True)
    preview_lines = st.number_input("미리보기용 문단 최대 개수", min_value=1, max_value=20, value=5)

# Upload area (main)
st.subheader("1) 워드 파일 업로드 (.docx)")
uploaded_file = st.file_uploader("학습 자료(.docx)를 업로드하세요", type=["docx"])

st.write("---")
st.subheader("간단 사용 설명")
st.markdown("""
- 파일 업로드 → 빈칸 비율 설정 → `시험지 생성 및 다운로드` 클릭  
- 생성된 파일에는 시험지(빈칸 처리된 본문)와 마지막 페이지의 `정답지`가 포함됩니다.
""")

# ---------- 도우미 함수들 ----------
TOKEN_RE = re.compile(r"\w+|[^\w\s]", re.UNICODE)  # 단어(또는 구두점) 단위 토큰화 대체 (nltk 토큰 대신)

def tokenize_preserve(text):
    """
    단순 토큰화: 단어/구두점을 분리. (NLTK와 유사하지만 join 시 공백을 넣어주는 방식이 다름)
    반환: 토큰 리스트
    """
    return TOKEN_RE.findall(text)

def is_candidate_token(tok, skip_punct=True):
    """빈칸 후보인지 판단 (구두점/숫자/심지어 한글 포함 모두 처리 가능)."""
    if skip_punct and re.fullmatch(r"[^\w\s]", tok):
        return False
    # 토큰에 최소 하나의 알파벳/한글/숫자가 있으면 후보로 본다
    return bool(re.search(r"[A-Za-z0-9\uac00-\ud7a3]", tok))

def assemble_from_tokens(tokens):
    """
    토큰 리스트를 문자열로 복원. 토큰화 방식에 따라 띄어쓰기 규칙을 단순화.
    (Punctuation 앞에는 공백을 제거)
    """
    text = ""
    for i, t in enumerate(tokens):
        if i == 0:
            text += t
            continue
        # 현재 토큰이 구두점이면 바로 붙이고, 아니면 앞에 공백 추가
        if re.fullmatch(r"[^\w\s]", t):
            text += t
        else:
            # 이전 토큰이 구두점이면 공백 없이 붙임 (예: "word," + "next" -> "word,next" is undesired)
            # 더 안전하게 항상 공백 추가
            text += " " + t
    return text

def generate_random_blanks_from_text(text, ratio, skip_punct=True):
    """
    text -> 토큰화 -> 후보 토큰 중 비율에 따라 무작위로 선택 -> 선택 토큰을 언더바로 대체
    반환: (masked_text_str, blanks_dict)
      blanks_dict: { token_global_index_in_tokens_list: original_token, ... }
    """
    tokens = tokenize_preserve(text)
    candidate_indices = [i for i, tok in enumerate(tokens) if is_candidate_token(tok, skip_punct=skip_punct)]
    n_blanks = max(1, int(len(candidate_indices) * ratio / 100)) if candidate_indices else 0
    blanks = {}
    if n_blanks > 0 and candidate_indices:
        chosen = random.sample(candidate_indices, min(n_blanks, len(candidate_indices)))
        for idx in chosen:
            blanks[idx] = tokens[idx]
            # 언더바 길이는 원래 토큰 길이만큼 (한글/영문 혼합도 처리)
            tokens[idx] = "_" * len(tokens[idx])
    masked = assemble_from_tokens(tokens)
    return masked, blanks

def set_paragraph_border(paragraph):
    """
    paragraph에 단순한 single border(테두리) 적용
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    # 기존에 pBdr가 있으면 제거(중복 방지)
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    pPr.append(pBdr)

def set_runs_font(paragraph, size_pt=12, bold=False):
    """
    paragraph의 모든 run에 대해 폰트 크기/볼드 설정
    """
    for run in paragraph.runs:
        run.font.size = Pt(size_pt)
        run.font.bold = bold

def set_cell_font(cell, size_pt=12, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size_pt)
            r.font.bold = bold

# ---------- 핵심: docx 생성 함수 ----------
def process_docx_with_answer(file_like, ratio, skip_punct=True):
    """
    업로드된 .docx 파일(파일 객체)을 받아
    - 빈칸 처리된 시험지(본문)
    - 마지막 페이지에 정답지
    를 포함한 새로운 Document를 생성하여 BytesIO로 반환
    """
    src = Document(file_like)
    dst = Document()

    # 여백 설정 (인치 단위)
    for section in dst.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --- 상단 시험지 정보 (테이블 형태) ---
    header_table = dst.add_table(rows=2, cols=4)
    header_table.style = 'Table Grid'
    header_table.autofit = True
    # 첫줄
    header_table.cell(0, 0).text = "반:"
    header_table.cell(0, 1).text = ""
    header_table.cell(0, 2).text = "이름:"
    header_table.cell(0, 3).text = ""
    # 둘째줄
    header_table.cell(1, 0).text = "점수:"
    header_table.cell(1, 1).text = ""
    header_table.cell(1, 2).text = "선생님 확인:"
    header_table.cell(1, 3).text = ""

    # 셀 폰트 조정
    for row in header_table.rows:
        for cell in row.cells:
            set_cell_font(cell, size_pt=11, bold=False)

    dst.add_paragraph("")  # 간격

    all_answers = []  # [{'para_index': n, 'original': text, 'blanks': {idx: token,...}} , ...]

    para_counter = 0
    for para in src.paragraphs:
        text = para.text.strip()
        if not text:
            # 빈 줄도 그대로 추가(공백)
            dst.add_paragraph("")
            continue

        masked, blanks = generate_random_blanks_from_text(text, ratio, skip_punct=skip_punct)
        p = dst.add_paragraph(masked)
        set_runs_font(p, size_pt=11, bold=False)
        # paragraph 테두리 적용
        set_paragraph_border(p)

        if blanks:
            all_answers.append({'index': para_counter, 'original': text, 'blanks': blanks})
        para_counter += 1

    # --- 답지 (마지막 페이지) ---
    dst.add_page_break()
    title = dst.add_paragraph("📝 정답지 (Answer Sheet)")
    # title runs에 폰트 적용
    set_runs_font(title, size_pt=13, bold=True)

    if not all_answers:
        dst.add_paragraph("빈칸으로 표시된 항목이 없습니다.")
    else:
        for i, item in enumerate(all_answers, start=1):
            # 정답 표시 형식: "1) [원문 일부] -> 답: token1, token2"
            sorted_idxs = sorted(item['blanks'].keys())
            answers = [item['blanks'][idx] for idx in sorted_idxs]
            answer_line = f"{i}. {item['original']}\n   답: " + ", ".join(answers)
            p = dst.add_paragraph(answer_line)
            set_runs_font(p, size_pt=11, bold=False)

    # 메모리에 저장
    out = BytesIO()
    dst.save(out)
    out.seek(0)
    return out

# ---------- UI 동작 ----------
if uploaded_file is not None:
    st.success("파일 업로드 확인됨")
    st.info("문서에서 랜덤으로 단어를 선택해 빈칸 처리합니다. (답지는 마지막 페이지에 추가됩니다.)")

    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("업로드된 문서 미리보기 (문단 일부)")
        try:
            doc_temp = Document(uploaded_file)
            lines = []
            for i, p in enumerate(doc_temp.paragraphs):
                if p.text.strip():
                    lines.append(p.text.strip())
                if len(lines) >= preview_lines:
                    break
            if lines:
                for i, l in enumerate(lines, 1):
                    st.write(f"**{i}.** {l}")
            else:
                st.write("문서에 미리보기할 텍스트가 없습니다.")
        except Exception as e:
            st.error("파일 미리보기 실패: 업로드된 파일이 올바른 .docx 파일인지 확인하세요.")
            st.write(e)

    with col2:
        st.subheader("생성 옵션")
        st.write(f"- 빈칸 비율: **{blank_ratio}%**")
        st.write(f"- 구두점 제외: **{keep_punct}**")
        st.write("")
        st.markdown("**다운로드**")
        if st.button("▶ 시험지 생성 및 다운로드"):
            try:
                # 업로드된 파일을 처음부터 다시 읽기 위해 .seek(0)
                uploaded_file.seek(0)
                output = process_docx_with_answer(uploaded_file, blank_ratio, skip_punct=keep_punct)
                st.success("시험지 생성 완료! 아래에서 다운로드하세요.")
                st.download_button(
                    label="⬇️ 시험지(.docx) 다운로드 (문제 + 답지 포함)",
                    data=output,
                    file_name="blank_test_with_answer.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error("시험지 생성 중 오류가 발생했습니다.")
                st.exception(e)
else:
    st.info("먼저 워드(.docx) 파일을 업로드하세요.")

st.markdown("---")
st.caption("Made with ❤️  ·  Your Blank Test Generator")
